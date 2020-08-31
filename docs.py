import twl
from docx import Document
from docx.enum.section import WD_SECTION

# didn't work as expected.
# def _set_columns(section, cols):
#     """
#     sets number of columns through xpath.
#     This is a workaround found on https://github.com/python-openxml/python-docx/issues/167
#     calls to _set_columns(section, ...) should be replaced by section.set_... after the issue is
#     fixed and set_columns or similar is added added to docx
#     """
#     section._sectPr.xpath("./w:cols")[0].set(_set_columns.WNS_COLS_NUM, str(cols))
#     # section.right_margin = Inches(0)


# _set_columns.WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


def fix_filename(fname, postfix='.docx'):
    if fname.lower().endswith(postfix):
        return fname

    return fname + postfix


def doc_by_letter_contains(wolf, cols, fname):
    doc = Document()
    full = twl.Wolf()

    for c in twl.letters:
        pup = wolf.starts(c)
        if len(pup) == 0:
            print('0', end='', flush=True)
            continue
        p = doc.add_paragraph()
        for word in pup.words:
            count = len(full.contains(word))
            p.add_run("{} {}\n".format(word, count))
        doc.add_paragraph('\n')
        print('.', end='', flush=True)

    print()

    if not fname:
        fname = str(wolf)

    doc.save(fix_filename(fname))


def doc_by_letter(wolf, cols, fname):
    doc = Document()

    for c in twl.letters:
        pup = wolf.starts(c)
        if len(pup) == 0:
            continue
        doc.add_paragraph(twl.prep_words(pup.words))
        doc.add_paragraph('\n')

    if not fname:
        fname = str(wolf)

    doc.save(fix_filename(fname))


wolf = twl.Wolf()


def byNumber(byletter=doc_by_letter, start=2, end=9):
    for n in range(start, end+1):
        print("preparing", n, "letters long")
        byletter(wolf.len(n), 9, "{}ltr".format(n))


def digraphs():
    doc = Document()

    pack = [(pair, wolf.contains(pair)) for pair in twl.digraphs]
    pack.sort(reverse=True, key=lambda p: len(p[1]))

    for pair, pup in pack:
        run = '{} {}'.format(pair, len(pup))
        doc.add_paragraph(run)
        print(run)

    doc.save(fix_filename('digraphs'))


if __name__ == "__main__":
    # byNumber()
    # digraphs()
    byNumber(doc_by_letter_contains, end=4)
    byNumber(start=5)
